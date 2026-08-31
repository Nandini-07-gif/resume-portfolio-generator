import os
import re
import json
import sys
import io
import base64
import zipfile
import xml.etree.ElementTree as ET
import urllib.parse
import webbrowser
from http.server import HTTPServer, ThreadingHTTPServer, BaseHTTPRequestHandler
from datetime import datetime
from dotenv import load_dotenv

# Prevent creation of __pycache__ bytecode files
sys.dont_write_bytecode = True

# Base directory of the script file
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Reconfigure stdout to utf-8 if possible for Windows terminal compatibility
if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

# Safe import for google.generativeai
try:
    import google.generativeai as genai
    HAS_GEMINI = True
except ImportError:
    genai = None
    HAS_GEMINI = False

def _extract_pdf_raw_strings(file_bytes: bytes) -> str:
    """Fallback parser to extract text strings from PDF byte streams."""
    try:
        raw = file_bytes.decode('latin-1', errors='ignore')
        matches = re.findall(r'\(([^()\r\n]{2,})\)\s*T[jJ]', raw)
        if not matches:
            matches = re.findall(r'\[\s*\(([^()\r\n]{2,})\)\s*\]\s*TJ', raw)
        if matches:
            cleaned_strings = [m.strip() for m in matches if len(m.strip()) > 1]
            return "\n".join(cleaned_strings)
    except Exception:
        pass
    return ""

def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract clean plain text from uploaded files (.txt, .pdf, .docx, .doc, .json)."""
    if not file_bytes:
        return ""
    
    ext = os.path.splitext(filename)[1].lower() if filename else ""

    # 1. JSON file
    if ext == '.json':
        try:
            raw_str = file_bytes.decode('utf-8', errors='ignore')
            data = json.loads(raw_str)
            if isinstance(data, dict):
                lines = []
                for k, v in data.items():
                    if isinstance(v, str) and v:
                        lines.append(f"{k.capitalize()}: {v}")
                    elif isinstance(v, list) and v:
                        lines.append(f"{k.capitalize()}:")
                        for item in v:
                            if isinstance(item, str):
                                lines.append(f"- {item}")
                            elif isinstance(item, dict):
                                item_str = " | ".join(str(val) for val in item.values() if val)
                                lines.append(f"- {item_str}")
                    elif isinstance(v, dict) and v:
                        lines.append(f"{k.capitalize()}:")
                        for sub_k, sub_v in v.items():
                            if sub_v:
                                lines.append(f"  {sub_k}: {sub_v}")
                if lines:
                    return "\n".join(lines)
                return raw_str
        except Exception:
            pass

    # 2. PDF file
    if ext == '.pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text).strip()
            if text:
                return text
        except Exception:
            pass

        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text).strip()
            if text:
                return text
        except Exception:
            pass

        raw_pdf_text = _extract_pdf_raw_strings(file_bytes)
        if raw_pdf_text:
            return raw_pdf_text

    # 3. DOCX / DOC file
    if ext in ('.docx', '.doc'):
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs).strip()
            if text:
                return text
        except Exception:
            pass

        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                xml_content = z.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                text_nodes = [node.text for node in tree.iter() if node.tag.endswith('}t') and node.text]
                text = " ".join(text_nodes).strip()
                if text:
                    return text
        except Exception:
            pass

    # 4. Fallback / Plain Text (.txt or unknown text file)
    for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
        try:
            text = file_bytes.decode(encoding).strip()
            if text:
                text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
                return text
        except UnicodeDecodeError:
            continue

    return ""

DEFAULT_PORTFOLIO_DATA = {
    "name": "",
    "headline": "",
    "professional_summary": "",
    "skills": [],
    "education": [],
    "experience": [],
    "projects": [],
    "achievements": [],
    "contact": {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "other_links": [],
    },
}

SAMPLE_PORTFOLIO_DATA = {
    "name": "Nandini Saraswat",
    "headline": "Computer Science Student | Web Development Intern",
    "professional_summary": "First-year Computer Science undergraduate specializing in AI, Analytics, and Full-Stack Web Development. Experienced in building Python backends, web applications, and predictive machine learning models.",
    "skills": ["Python", "Java", "SQL", "HTML/CSS", "JavaScript", "Flask", "React", "Vite", "Git", "MySQL", "Gemini API", "Public Speaking", "Team Coordination", "Event Management"],
    "education": [{"degree": "Bachelor of Technology in Computer Science (AI & Analytics)", "institution": "GLA University, Mathura", "duration": "2025 - Present"}],
    "experience": [{"role": "Web Development Intern", "company": "Infosys Prodigy", "duration": "2025", "responsibilities": ["Developed responsive web interfaces and assisted with backend features.", "Collaborated on web application features using modern HTML, CSS, and JavaScript."]}],
    "projects": [
        {"title": "Zenith 2", "description": "Project Management & Collaboration Platform built using Python and Flask.", "technologies": ["Python", "Flask", "SQL"]},
        {"title": "AI Insurance Claim Verification Agent", "description": "Automated verification system using public registries to validate claim data.", "technologies": ["Python", "Gemini API"]}
    ],
    "achievements": ["Overall Coordinator, College Hackathon (2025)", "Founder / Team Lead, Aashayein Social Cause Club"],
    "contact": {"email": "contact@example.com", "phone": "", "linkedin": "linkedin.com/in/example", "github": "github.com/example", "other_links": []}
}

THEME_MAP = {
    "1": ("theme-glassmorphism", "Glassmorphism Creative"),
    "2": ("theme-cyberpunk", "White Modern Editorial"),
    "3": ("theme-terminal", "Developer Terminal"),
    "4": ("theme-minimalist", "Modern Minimalist"),
}

def clean_resume_text(raw_text: str) -> str:
    """Cleans raw text by removing extra spaces and weird characters."""
    cleaned_text = re.sub(r'[\r\t]', ' ', raw_text)
    cleaned_text = re.sub(r' +', ' ', cleaned_text)
    cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
    return cleaned_text.strip()

SECTION_KEY_MAP = {
    "summary": "summary",
    "professional summary": "summary",
    "skills": "skills",
    "education": "education",
    "experience": "experience",
    "work experience": "experience",
    "projects": "projects",
    "awards & leadership": "achievements",
    "awards": "achievements",
    "achievements": "achievements",
}

def _has_meaningful_data(data: dict) -> bool:
    """Returns True if parsed data contains usable portfolio content."""
    if data.get("name") or data.get("headline") or data.get("professional_summary"):
        return True
    if data.get("skills") or data.get("experience") or data.get("projects"):
        return True
    if data.get("education") or data.get("achievements"):
        return True
    contact = data.get("contact", {})
    if isinstance(contact, dict) and any(contact.values()):
        return True
    return False

def _extract_contact(line: str) -> dict:
    """Pull email, phone, LinkedIn, and GitHub from a contact line."""
    contact = {"email": "", "phone": "", "linkedin": "", "github": "", "other_links": []}
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", line, re.I)
    if email_match:
        contact["email"] = email_match.group(0)
    phone_match = re.search(r"(?:\+?\d{1,3}[\s-]?)?\(?\d{2,4}\)?[\s-]?\d{3,4}[\s-]?\d{3,4}", line)
    if phone_match:
        contact["phone"] = phone_match.group(0).strip()
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", line, re.I)
    if linkedin_match:
        contact["linkedin"] = linkedin_match.group(0).rstrip("|,")
    github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/\S+", line, re.I)
    if github_match:
        contact["github"] = github_match.group(0).rstrip("|,")
    return contact

def _split_resume_sections(text: str) -> tuple[list[str], dict[str, list[str]]]:
    """Split resume into preamble lines (before first header) and named sections."""
    preamble = []
    sections = {}
    current_key = None

    for line in text.split("\n"):
        stripped = line.strip()
        header_key = SECTION_KEY_MAP.get(stripped.lower()) if stripped else None

        if header_key:
            current_key = header_key
            sections.setdefault(current_key, [])
        elif current_key:
            sections.setdefault(current_key, []).append(stripped)
        elif stripped:
            preamble.append(stripped)

    return preamble, sections

def _parse_skills(lines: list[str]) -> list[str]:
    skills = []
    for line in lines:
        if not line:
            continue
        cleaned = re.sub(r"^[-•*]\s*", "", line).strip()
        if not cleaned:
            continue
        if ":" in cleaned:
            _, items_part = cleaned.split(":", 1)
            for item in re.split(r"[,;|]", items_part):
                skill = item.strip()
                if skill:
                    skills.append(skill)
        elif "," in cleaned:
            for item in cleaned.split(","):
                skill = item.strip()
                if skill:
                    skills.append(skill)
        else:
            skills.append(cleaned)
    return skills

def _parse_education(lines: list[str]) -> list[dict]:
    entries = []
    non_empty = [line for line in lines if line]
    i = 0
    while i < len(non_empty):
        institution = non_empty[i]
        degree = ""
        duration = ""
        if i + 1 < len(non_empty):
            next_line = non_empty[i + 1]
            if "|" in next_line:
                parts = [part.strip() for part in next_line.split("|")]
                degree = parts[0]
                duration = parts[1] if len(parts) > 1 else ""
            else:
                degree = next_line
            i += 2
        else:
            i += 1
        entries.append({"degree": degree, "institution": institution, "duration": duration})
    return entries

def _parse_role_block(lines: list[str], default_role_key: str, default_company_key: str) -> list[dict]:
    entries = []
    current = None

    for line in lines:
        if not line:
            continue
        if re.match(r"^[-•*]\s+", line):
            if current:
                current["responsibilities"].append(re.sub(r"^[-•*]\s+", "", line).strip())
            continue

        if current:
            entries.append(current)

        if "—" in line:
            parts = [part.strip() for part in line.split("—", 1)]
            current = {
                default_company_key: parts[0],
                default_role_key: parts[1],
                "duration": "",
                "responsibilities": [],
            }
        elif " - " in line and not line.startswith("-"):
            parts = [part.strip() for part in line.split(" - ", 1)]
            current = {
                default_company_key: parts[0],
                default_role_key: parts[1],
                "duration": "",
                "responsibilities": [],
            }
        else:
            current = {
                default_role_key: line.strip(),
                default_company_key: "",
                "duration": "",
                "responsibilities": [],
            }

    if current:
        entries.append(current)

    return entries

def _parse_experience(lines: list[str]) -> list[dict]:
    raw_entries = _parse_role_block(lines, "role", "company")
    experience = []
    for entry in raw_entries:
        experience.append({
            "role": entry.get("role", ""),
            "company": entry.get("company", ""),
            "duration": entry.get("duration", ""),
            "responsibilities": entry.get("responsibilities", []),
        })
    return experience

def _parse_projects(lines: list[str]) -> list[dict]:
    raw_entries = _parse_role_block(lines, "title", "company")
    projects = []
    for entry in raw_entries:
        description = " ".join(entry.get("responsibilities", []))
        projects.append({
            "title": entry.get("title", ""),
            "description": description,
            "technologies": [],
        })
    return projects

def _parse_achievements(lines: list[str]) -> list[str]:
    achievements = []
    for line in lines:
        if not line:
            continue
        cleaned = re.sub(r"^[-•*]\s*", "", line).strip()
        if cleaned:
            achievements.append(cleaned)
    return achievements

def parse_resume_locally(text: str) -> dict:
    """Parse plain-text resume into structured portfolio data without AI."""
    data = DEFAULT_PORTFOLIO_DATA.copy()
    if not text.strip():
        return data

    preamble, sections = _split_resume_sections(text)

    if preamble:
        data["name"] = preamble[0]
        contact_line = ""
        headline = ""

        for line in preamble[1:]:
            extracted = _extract_contact(line)
            if any(extracted.values()):
                contact_line = line
                data["contact"] = extracted
            elif not headline:
                headline = line

        if headline:
            data["headline"] = headline
        if contact_line and not any(data["contact"].values()):
            data["contact"] = _extract_contact(contact_line)

    if "summary" in sections:
        data["professional_summary"] = " ".join(line for line in sections["summary"] if line).strip()

    if "skills" in sections:
        data["skills"] = _parse_skills(sections["skills"])

    if "education" in sections:
        data["education"] = _parse_education(sections["education"])

    if "experience" in sections:
        data["experience"] = _parse_experience(sections["experience"])

    if "projects" in sections:
        data["projects"] = _parse_projects(sections["projects"])

    if "achievements" in sections:
        data["achievements"] = _parse_achievements(sections["achievements"])

    return data

def parse_structured_input(text: str) -> dict:
    """
    Parse the __STRUCTURED__ format sent by the new field-based UI.
    Each line is 'FieldName: value'. Multi-line sections use newlines within the value.
    Returns a fully populated portfolio dict, bypassing AI/local parsing entirely.
    """
    if not text.startswith('__STRUCTURED__'):
        return None  # Not our format

    data = DEFAULT_PORTFOLIO_DATA.copy()
    # Build a dict from Key: value lines
    fields = {}
    current_key = None
    for line in text.split('\n'):
        if line == '__STRUCTURED__':
            continue
        # Detect "Key: value" pattern (key has no spaces except in known multi-word keys)
        m = re.match(r'^(Name|Title|Summary|Skills|Experience|Education|Projects|Achievements|Email|LinkedIn|GitHub):\s*(.*)', line)
        if m:
            current_key = m.group(1)
            fields[current_key] = m.group(2).strip()
        elif current_key and line.strip():
            # Continuation line for multi-line fields
            fields[current_key] = fields.get(current_key, '') + '\n' + line.strip()

    # Map fields → portfolio data
    if fields.get('Name'):
        data['name'] = fields['Name']
    if fields.get('Title'):
        data['headline'] = fields['Title']
    if fields.get('Summary'):
        data['professional_summary'] = fields['Summary']

    # Skills: comma-separated
    if fields.get('Skills'):
        data['skills'] = [s.strip() for s in fields['Skills'].split(',') if s.strip()]

    # Experience: each line = "Role | Company | Year | Description"
    if fields.get('Experience'):
        exp_list = []
        for line in fields['Experience'].split('\n'):
            line = line.strip()
            if not line:
                continue
            parts = [p.strip() for p in line.split('|')]
            exp_list.append({
                'role': parts[0] if len(parts) > 0 else line,
                'company': parts[1] if len(parts) > 1 else '',
                'duration': parts[2] if len(parts) > 2 else '',
                'responsibilities': [parts[3]] if len(parts) > 3 else [],
            })
        data['experience'] = exp_list

    # Education: each line = "Degree | Institution | Year"
    if fields.get('Education'):
        edu_list = []
        for line in fields['Education'].split('\n'):
            line = line.strip()
            if not line:
                continue
            parts = [p.strip() for p in line.split('|')]
            edu_list.append({
                'degree': parts[0] if len(parts) > 0 else line,
                'institution': parts[1] if len(parts) > 1 else '',
                'duration': parts[2] if len(parts) > 2 else '',
            })
        data['education'] = edu_list

    # Projects: each line = "Name | Description | Tech1, Tech2"
    if fields.get('Projects'):
        proj_list = []
        for line in fields['Projects'].split('\n'):
            line = line.strip()
            if not line:
                continue
            parts = [p.strip() for p in line.split('|')]
            techs = [t.strip() for t in parts[2].split(',')] if len(parts) > 2 else []
            proj_list.append({
                'title': parts[0] if len(parts) > 0 else line,
                'description': parts[1] if len(parts) > 1 else '',
                'technologies': techs,
            })
        data['projects'] = proj_list

    # Achievements: one per line
    if fields.get('Achievements'):
        data['achievements'] = [
            line.strip() for line in fields['Achievements'].split('\n') if line.strip()
        ]

    # Contact
    contact = data.get('contact', {})
    if fields.get('Email'):
        contact['email'] = fields['Email']
    if fields.get('LinkedIn'):
        lk = fields['LinkedIn']
        contact['linkedin'] = lk if lk.startswith('http') else 'https://' + lk
    if fields.get('GitHub'):
        gh = fields['GitHub']
        contact['github'] = gh if gh.startswith('http') else 'https://' + gh
    data['contact'] = contact

    return data

def configure_gemini():
    """Loads the API key and sets up the Gemini model."""
    if not HAS_GEMINI:
        return None
    env_path = os.path.join(BASE_DIR, ".env")
    load_dotenv(dotenv_path=env_path)
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return None
    try:
        genai.configure(api_key=api_key)
        return genai.GenerativeModel("gemini-2.5-flash")
    except Exception:
        return None

def build_prompt(cleaned_text: str) -> str:
    """Builds the strict instruction prompt sent to Gemini."""
    schema = """
    {
        "name": "",
        "headline": "",
        "professional_summary": "",
        "skills": [],
        "education": [{"degree": "", "institution": "", "duration": ""}],
        "experience": [{"role": "", "company": "", "duration": "", "responsibilities": []}],
        "projects": [{"title": "", "description": "", "technologies": []}],
        "achievements": [],
        "contact": {"email": "", "phone": "", "linkedin": "", "github": "", "other_links": []}
    }
    """

    return f"""You are extracting structured portfolio data from a resume.
        Rules:
        - Use ONLY information explicitly present in the resume text below.
        - Do NOT invent or assume any skills, experience, projects, companies, dates, achievements, or links.
        - If information is not present, use an empty string "" or empty list [].
        - Keep the professional summary concise (2-3 sentences) and strictly factual.
        - Respond with VALID JSON ONLY. No markdown code fences, no explanations, no extra text.

        Return JSON matching exactly this structure:
        {schema}

        ---RESUME START---
        {cleaned_text}
        ---RESUME END---
        """

def call_gemini(model, prompt: str) -> str:
    """Sends the prompt to Gemini and returns the raw text response."""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"[ERROR] Could not reach Gemini API. Details: {e}")
        return ""

def parse_and_validate_json(raw_response: str) -> dict:
    """Cleans raw API response, converts JSON to Python dictionary safely, and handles missing values."""
    if not raw_response:
        return DEFAULT_PORTFOLIO_DATA.copy()

    cleaned = raw_response.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[-1]
    if cleaned.endswith("```"):
        cleaned = cleaned.rsplit("```", 1)[0]
    cleaned = cleaned.strip()

    try:
        data = json.loads(cleaned)
        if not isinstance(data, dict):
            return DEFAULT_PORTFOLIO_DATA.copy()
    except (json.JSONDecodeError, TypeError):
        return DEFAULT_PORTFOLIO_DATA.copy()

    validated_data = DEFAULT_PORTFOLIO_DATA.copy()
    for key, default_val in DEFAULT_PORTFOLIO_DATA.items():
        val = data.get(key)
        if val is not None:
            validated_data[key] = val

    contact_data = validated_data.get("contact", {})
    if not isinstance(contact_data, dict):
        contact_data = {}

    validated_data["contact"] = {
        "email": contact_data.get("email", ""),
        "phone": contact_data.get("phone", ""),
        "linkedin": contact_data.get("linkedin", ""),
        "github": contact_data.get("github", ""),
        "other_links": contact_data.get("other_links", []),
    }

    return validated_data

def remove_section(html_content, section_id):
    """Remove an entire section from HTML if it has no content"""
    pattern = r'<section id="' + section_id + r'">.*?</section>'
    return re.sub(pattern, '', html_content, flags=re.DOTALL)

def generate_html(data, theme_class="theme-minimalist"):
    """Generate portfolio.html from JSON data using template.html and chosen theme_class"""
    try:
        template_file = os.path.join(BASE_DIR, 'template.html')
        with open(template_file, 'r', encoding='utf-8') as file:
            html_template = file.read()
        
        html_content = html_template.replace('{THEME_CLASS}', theme_class)
        heading = data.get('heading', data.get('headline', 'Professional'))
        
        html_content = html_content.replace('{NAME}', data.get('name', 'Not Specified'))
        html_content = html_content.replace('{HEADING}', heading)
        html_content = html_content.replace('{SUMMARY}', data.get('professional_summary', data.get('summary', '')))
        html_content = html_content.replace('{DATE}', datetime.now().strftime('%B %d, %Y'))
        
        # 1. Contact Section
        contact = data.get('contact', {})
        if contact and any(contact.values()):
            contact_html = ''
            if contact.get('email'):
                contact_html += f'<div class="contact-item"><i class="fas fa-envelope"></i><a href="mailto:{contact["email"]}">{contact["email"]}</a></div>'
            if contact.get('phone'):
                contact_html += f'<div class="contact-item"><i class="fas fa-phone"></i><a href="tel:{contact["phone"]}">{contact["phone"]}</a></div>'
            if contact.get('linkedin'):
                l_url = contact['linkedin'] if contact['linkedin'].startswith('http') else 'https://' + contact['linkedin']
                contact_html += f'<div class="contact-item"><i class="fab fa-linkedin"></i><a href="{l_url}" target="_blank">LinkedIn</a></div>'
            if contact.get('github'):
                g_url = contact['github'] if contact['github'].startswith('http') else 'https://' + contact['github']
                contact_html += f'<div class="contact-item"><i class="fab fa-github"></i><a href="{g_url}" target="_blank">GitHub</a></div>'
            html_content = html_content.replace('{CONTACT_CONTENT}', contact_html)
        else:
            html_content = remove_section(html_content, 'contact')
        
        # 2. Skills Section
        skills_list = data.get('skills', [])
        if skills_list:
            skills_html = ''.join([f'<span class="skill-tag">{s}</span>\n' for s in skills_list])
            html_content = html_content.replace('{SKILLS_CONTENT}', skills_html)
        else:
            html_content = remove_section(html_content, 'skills')
        
        # 3. Experience Section
        experience_list = data.get('experience', [])
        if experience_list:
            exp_html = ''
            for exp in experience_list:
                role = exp.get('role', exp.get('title', ''))
                company = exp.get('company', '')
                resps = exp.get('responsibilities', [])
                resp_text = ' | '.join(resps) if isinstance(resps, list) and resps else exp.get('description', '')
                exp_html += f'<div class="experience-item"><h3>{role} at {company}</h3><p class="date">{exp.get("duration", "")}</p><p class="description">{resp_text}</p></div>\n'
            html_content = html_content.replace('{EXPERIENCE_CONTENT}', exp_html)
        else:
            html_content = remove_section(html_content, 'experience')
        
        # 4. Education Section
        education_list = data.get('education', [])
        if education_list:
            edu_html = ''
            for edu in education_list:
                edu_html += f'<div class="education-item"><h3>{edu.get("degree", "")}</h3><p class="institution">{edu.get("institution", "")}</p><p class="date">{edu.get("duration", "")}</p></div>\n'
            html_content = html_content.replace('{EDUCATION_CONTENT}', edu_html)
        else:
            html_content = remove_section(html_content, 'education')
        
        # 5. Projects Section
        projects_list = data.get('projects', [])
        if projects_list:
            proj_html = ''
            for proj in projects_list:
                name = proj.get('title', proj.get('name', ''))
                desc = proj.get('description', '')
                techs = proj.get('technologies', [])
                t_html = '<div class="technologies">' + ''.join([f'<span class="tech-tag">{t}</span>' for t in techs]) + '</div>' if techs else ''
                proj_html += f'<div class="project-item"><h3>{name}</h3><p class="description">{desc}</p>{t_html}</div>\n'
            html_content = html_content.replace('{PROJECTS_CONTENT}', proj_html)
        else:
            html_content = remove_section(html_content, 'projects')
        
        # 6. Achievements Section
        achievements_list = data.get('achievements', [])
        if achievements_list:
            ach_html = ''
            for ach in achievements_list:
                if isinstance(ach, str):
                    ach_html += f'<div class="achievement-item"><h3>🏅 {ach}</h3></div>\n'
                elif isinstance(ach, dict):
                    ach_html += f'<div class="achievement-item"><h3>{ach.get("title","")}</h3><p class="description">{ach.get("description","")}</p><p class="date">{ach.get("year","")}</p></div>\n'
            html_content = html_content.replace('{ACHIEVEMENTS_CONTENT}', ach_html)
        else:
            html_content = remove_section(html_content, 'achievements')
        
        # Clean remaining placeholders & whitespace
        for placeholder in ['{CONTACT_CONTENT}', '{SKILLS_CONTENT}', '{EXPERIENCE_CONTENT}', '{EDUCATION_CONTENT}', '{PROJECTS_CONTENT}', '{ACHIEVEMENTS_CONTENT}']:
            html_content = html_content.replace(placeholder, '')
        html_content = re.sub(r'\n\s*\n', '\n', html_content)
        
        output_file = os.path.join(BASE_DIR, 'portfolio.html')
        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(html_content)
        
        return True
    except Exception as e:
        print(f"[ERROR] HTML Generation failed: {e}")
        return False

# ==========================================================
# WEB APPLICATION DASHBOARD HTML
# ==========================================================
WEB_APP_HTML = """<!DOCTYPE html>
<html lang="en" data-theme="light">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Portfolio Studio — Resume to Web Portfolio Generator</title>
    <meta name="description" content="Convert resume details into an elegant, single-page web portfolio. Clean, responsive, and customizable.">
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@500;600;700;800&family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/css/all.min.css">
    <style>
        /* =============================================
           DESIGN SYSTEM & THEME TOKENS (LIGHT DEFAULT)
           ============================================= */
        :root, [data-theme="light"] {
            --bg-base:       #faf9f6;
            --bg-surface:    #ffffff;
            --bg-raised:     #f1f5f9;
            --bg-overlay:    #ffffff;
            --border:        #e2e8f0;
            --border-hover:  #cbd5e1;
            --accent:        #4f46e5;
            --accent-hover:  #4338ca;
            --accent-subtle: #e0e7ff;
            --accent-green:  #059669;
            --accent-amber:  #d97706;
            --text-primary:  #0f172a;
            --text-secondary:#475569;
            --text-muted:    #64748b;
            --radius-sm:     8px;
            --radius-md:     12px;
            --radius-lg:     16px;
            --radius-xl:     24px;
            --shadow-sm:     0 1px 3px rgba(15,23,42,0.06);
            --shadow-md:     0 6px 20px -4px rgba(15,23,42,0.08);
            --shadow-lg:     0 16px 40px -8px rgba(15,23,42,0.12);
            --transition:    0.2s ease;
        }

        [data-theme="dark"] {
            --bg-base:       #0b0f19;
            --bg-surface:    #151c2c;
            --bg-raised:     #1e293b;
            --bg-overlay:    #1e293b;
            --border:        #334155;
            --border-hover:  #475569;
            --accent:        #6366f1;
            --accent-hover:  #818cf8;
            --accent-subtle: rgba(99,102,241,0.15);
            --accent-green:  #10b981;
            --accent-amber:  #f59e0b;
            --text-primary:  #f8fafc;
            --text-secondary:#94a3b8;
            --text-muted:    #64748b;
            --shadow-sm:     0 1px 3px rgba(0,0,0,0.3);
            --shadow-md:     0 6px 20px rgba(0,0,0,0.4);
            --shadow-lg:     0 16px 40px rgba(0,0,0,0.6);
        }

        /* Dark mode — explicit readable overrides */
        [data-theme="dark"] .field-input {
            background: #1e293b;
            color: #f8fafc;
            border-color: #334155;
        }
        [data-theme="dark"] .field-input::placeholder { color: #64748b; }
        [data-theme="dark"] .btn-secondary {
            background: #1e293b;
            color: #e2e8f0;
            border-color: #334155;
        }
        [data-theme="dark"] .btn-secondary:hover { background: #273549; border-color: #475569; }
        [data-theme="dark"] .score-card { background: #1e293b; border-color: #334155; }
        [data-theme="dark"] .score-bar-bg { background: #334155; }
        [data-theme="dark"] .status-bar.error { color: #fca5a5; background: rgba(220,38,38,0.15); border-color: rgba(220,38,38,0.3); }
        [data-theme="dark"] .status-bar.info { color: #a5b4fc; }
        [data-theme="dark"] .status-bar.success { color: #6ee7b7; background: rgba(5,150,105,0.12); }
        [data-theme="dark"] .dropdown-menu { background: #1e293b; border-color: #334155; }
        [data-theme="dark"] .dropdown-item { color: #e2e8f0; border-color: #334155; }
        [data-theme="dark"] .dropdown-item:hover { background: rgba(99,102,241,0.15); }
        [data-theme="dark"] .modal-card { background: #151c2c; border-color: #334155; }
        [data-theme="dark"] .modal-option-btn { background: #1e293b; border-color: #334155; color: #e2e8f0; }
        [data-theme="dark"] .theme-card { background: #1e293b; border-color: #334155; }
        [data-theme="dark"] .theme-card h4 { color: #f8fafc; }
        [data-theme="dark"] .theme-card.active { background: rgba(99,102,241,0.2); border-color: #6366f1; }
        [data-theme="dark"] .preview-placeholder { color: #64748b; }
        [data-theme="dark"] .hero-sub { color: #94a3b8; }
        [data-theme="dark"] .hero h1 { color: #f8fafc; }
        [data-theme="dark"] .brand-name { color: #f8fafc; }
        [data-theme="dark"] .field-label { color: #e2e8f0; }
        [data-theme="dark"] .field-hint-text { color: #64748b; }
        [data-theme="dark"] .panel-title { color: #e2e8f0; }
        [data-theme="dark"] .feature-chip { color: #94a3b8; }
        [data-theme="dark"] .score-title { color: #e2e8f0; }
        [data-theme="dark"] .tab-btn { color: #94a3b8; }
        [data-theme="dark"] .tab-btn.active { color: #818cf8; background: #1e293b; }
        [data-theme="dark"] .tab-btn:hover { color: #f8fafc; background: #273549; }

        *, *::before, *::after { margin:0; padding:0; box-sizing:border-box; }
        html { scroll-behavior: smooth; }
        body {
            font-family: 'Inter', -apple-system, sans-serif;
            background: var(--bg-base);
            color: var(--text-primary);
            min-height: 100vh;
            line-height: 1.6;
            -webkit-font-smoothing: antialiased;
            transition: background-color 0.3s ease, color 0.3s ease;
        }

        h1, h2, h3, h4, .brand-name {
            font-family: 'Plus Jakarta Sans', sans-serif;
        }

        /* Architectural Grid Substrate */
        .bg-grid {
            position: fixed; inset: 0; z-index: 0;
            background-image:
                linear-gradient(var(--border) 1px, transparent 1px),
                linear-gradient(90deg, var(--border) 1px, transparent 1px);
            background-size: 48px 48px;
            opacity: 0.35;
            pointer-events: none;
        }

        .app-root { position: relative; z-index: 1; min-height: 100vh; display: flex; flex-direction: column; }

        /* =============================================
           WELCOME MODAL
           ============================================= */
        .modal-overlay {
            position: fixed; inset: 0; z-index: 999;
            background: rgba(15,23,42,0.5);
            backdrop-filter: blur(10px); -webkit-backdrop-filter: blur(10px);
            display: flex; align-items: center; justify-content: center;
            padding: 20px;
            opacity: 0; pointer-events: none;
            transition: opacity var(--transition);
        }
        .modal-overlay.active { opacity: 1; pointer-events: auto; }
        .modal-card {
            background: var(--bg-surface);
            border: 1px solid var(--border);
            border-radius: var(--radius-xl);
            max-width: 540px; width: 100%;
            padding: 36px;
            box-shadow: var(--shadow-lg);
            transform: translateY(16px) scale(0.98);
            transition: transform var(--transition);
            text-align: center;
        }
        .modal-overlay.active .modal-card { transform: translateY(0) scale(1); }
        .modal-icon {
            width: 56px; height: 56px; margin: 0 auto 18px;
            background: var(--accent-subtle);
            color: var(--accent);
            border-radius: var(--radius-lg);
            display: flex; align-items: center; justify-content: center;
            font-size: 24px;
        }
        .modal-card h2 { font-size: 1.6em; font-weight: 800; margin-bottom: 8px; color: var(--text-primary); }
        .modal-card p { color: var(--text-secondary); font-size: 0.95em; line-height: 1.6; margin-bottom: 24px; }
        .modal-options { display: grid; grid-template-columns: 1fr 1fr; gap: 14px; margin-bottom: 18px; }
        @media (max-width: 540px) { .modal-options { grid-template-columns: 1fr; } }
        .modal-option-btn {
            background: var(--bg-base); border: 1px solid var(--border);
            border-radius: var(--radius-md); padding: 18px 14px; cursor: pointer;
            text-align: center; transition: all var(--transition);
        }
        .modal-option-btn:hover { border-color: var(--accent); background: var(--accent-subtle); transform: translateY(-2px); }
        .modal-option-btn i { font-size: 1.5em; margin-bottom: 8px; display: block; color: var(--accent); }
        .modal-option-btn .opt-title { font-weight: 700; font-size: 0.9em; color: var(--text-primary); margin-bottom: 4px; }
        .modal-option-btn .opt-desc { font-size: 0.78em; color: var(--text-muted); line-height: 1.4; }
        .modal-close-link { font-size: 0.82em; color: var(--text-muted); cursor: pointer; text-decoration: underline; }
        .modal-close-link:hover { color: var(--text-primary); }

        /* =============================================
           NAVBAR / HEADER
           ============================================= */
        .topbar {
            position: sticky; top: 0; z-index: 100;
            display: flex; align-items: center; justify-content: space-between;
            padding: 0 32px; height: 64px;
            background: var(--bg-surface);
            border-bottom: 1px solid var(--border);
            box-shadow: var(--shadow-sm);
        }
        .topbar-brand { display: flex; align-items: center; gap: 10px; text-decoration: none; }
        .brand-icon {
            width: 32px; height: 32px;
            background: var(--accent);
            border-radius: var(--radius-sm);
            display: flex; align-items: center; justify-content: center;
            font-size: 16px; color: #fff; font-weight: 800;
        }
        .brand-name { font-size: 1.1em; font-weight: 800; color: var(--text-primary); letter-spacing: -0.02em; }
        .brand-name span { color: var(--accent); }
        .topbar-right { display: flex; align-items: center; gap: 10px; }

        .save-indicator {
            font-size: 0.78em; font-weight: 600; color: var(--accent-green);
            background: rgba(5,150,105,0.08); border: 1px solid rgba(5,150,105,0.2);
            padding: 4px 12px; border-radius: 20px; display: flex; align-items: center; gap: 6px;
        }

        /* =============================================
           HERO HEADER
           ============================================= */
        .hero { text-align: center; padding: 40px 24px 28px; max-width: 760px; margin: 0 auto; }
        .hero h1 { font-size: clamp(2em, 4vw, 2.8em); font-weight: 800; letter-spacing: -0.03em; line-height: 1.2; margin-bottom: 10px; color: var(--text-primary); }
        .hero h1 .accent-text { color: var(--accent); }
        .hero-sub { font-size: 1em; color: var(--text-secondary); max-width: 580px; margin: 0 auto 20px; line-height: 1.6; }

        /* =============================================
           MAIN WORKSPACE LAYOUT
           ============================================= */
        .workspace {
            max-width: 1440px; margin: 0 auto;
            padding: 0 28px 60px;
            display: grid; grid-template-columns: 600px 1fr; gap: 28px;
            align-items: start;
        }
        @media (max-width: 1180px) { .workspace { grid-template-columns: 1fr; } }

        /* =============================================
           PANELS & CARDS
           ============================================= */
        .panel {
            background: var(--bg-surface);
            border: 1px solid var(--border);
            border-radius: var(--radius-lg);
            overflow: hidden; box-shadow: var(--shadow-md);
        }
        .panel-header {
            display: flex; align-items: center; justify-content: space-between;
            padding: 16px 24px; border-bottom: 1px solid var(--border);
            background: var(--bg-raised);
        }
        .panel-title { display: flex; align-items: center; gap: 8px; font-size: 0.9em; font-weight: 700; color: var(--text-primary); }
        .panel-title i { color: var(--accent); }
        .panel-body { padding: 24px; }

        /* =============================================
           SCORE & COMPLETENESS BAR
           ============================================= */
        .score-card {
            background: var(--bg-raised); border: 1px solid var(--border);
            border-radius: var(--radius-md); padding: 14px 18px; margin-bottom: 20px;
            display: flex; align-items: center; justify-content: space-between; gap: 14px;
        }
        .score-info { flex: 1; }
        .score-title { font-size: 0.82em; font-weight: 700; color: var(--text-primary); margin-bottom: 6px; display: flex; align-items: center; justify-content: space-between; }
        .score-bar-bg { height: 6px; background: var(--border); border-radius: 3px; overflow: hidden; }
        .score-bar-fill { height: 100%; width: 25%; background: var(--accent); transition: width 0.3s ease; border-radius: 3px; }

        /* =============================================
           FORM TAB NAVIGATOR
           ============================================= */
        .form-tabs {
            display: flex; gap: 4px; background: var(--bg-raised);
            padding: 4px; border-radius: var(--radius-md); border: 1px solid var(--border);
            margin-bottom: 20px; overflow-x: auto;
        }
        .tab-btn {
            flex: 1; padding: 9px 12px; border-radius: var(--radius-sm);
            border: none; background: transparent; color: var(--text-secondary);
            font-size: 0.8em; font-weight: 600; cursor: pointer;
            display: flex; align-items: center; justify-content: center; gap: 6px;
            white-space: nowrap; transition: all var(--transition);
        }
        .tab-btn:hover { color: var(--text-primary); background: var(--bg-surface); }
        .tab-btn.active { background: var(--bg-surface); color: var(--accent); box-shadow: var(--shadow-sm); border: 1px solid var(--border); }

        .tab-content { display: none; }
        .tab-content.active { display: block; }

        /* =============================================
           FORM FIELDS
           ============================================= */
        .field-group { margin-bottom: 18px; }
        .field-label-row { display: flex; align-items: center; justify-content: space-between; margin-bottom: 6px; }
        .field-label { font-size: 0.8em; font-weight: 700; color: var(--text-primary); letter-spacing: 0.02em; display: flex; align-items: center; gap: 6px; }
        .field-label i { color: var(--accent); font-size: 0.9em; }
        .btn-ai-enhance {
            font-size: 0.72em; font-weight: 600; color: var(--accent);
            background: var(--accent-subtle); border: 1px solid rgba(79,70,229,0.2);
            padding: 3px 10px; border-radius: 12px; cursor: pointer;
            transition: all var(--transition); display: inline-flex; align-items: center; gap: 4px;
        }
        .btn-ai-enhance:hover { background: var(--accent); color: #fff; }

        .field-input {
            width: 100%; background: var(--bg-surface);
            border: 1px solid var(--border); border-radius: var(--radius-md);
            padding: 12px 16px; color: var(--text-primary);
            font-family: 'Inter', sans-serif; font-size: 0.9em; line-height: 1.5;
            outline: none; transition: border-color var(--transition), box-shadow var(--transition);
        }
        .field-input:focus { border-color: var(--accent); box-shadow: 0 0 0 3px var(--accent-subtle); }
        .field-input::placeholder { color: var(--text-muted); }
        textarea.field-input { resize: vertical; min-height: 100px; }
        .field-hint-text { font-size: 0.76em; color: var(--text-muted); margin-top: 4px; }

        /* =============================================
           BUTTONS & ACTIONS
           ============================================= */
        .btn {
            display: inline-flex; align-items: center; justify-content: center; gap: 8px;
            padding: 10px 18px; border-radius: var(--radius-md);
            font-family: 'Inter', sans-serif; font-size: 0.88em; font-weight: 600;
            border: none; cursor: pointer; text-decoration: none;
            transition: all var(--transition); white-space: nowrap; user-select: none;
        }
        .btn:hover { transform: translateY(-1px); opacity: 0.96; }
        .btn:active { transform: translateY(0); }
        .btn-primary { background: var(--accent); color: #fff; box-shadow: 0 4px 14px rgba(79,70,229,0.25); }
        .btn-primary:hover { background: var(--accent-hover); }
        .btn-secondary { background: var(--bg-raised); color: var(--text-primary); border: 1px solid var(--border); }
        .btn-secondary:hover { border-color: var(--border-hover); background: var(--bg-surface); }
        .btn-amber { background: var(--accent-amber); color: #fff; }
        .btn-green { background: var(--accent-green); color: #fff; }
        .btn-sm { padding: 7px 12px; font-size: 0.8em; }
        .btn-full { width: 100%; padding: 14px; font-size: 0.95em; }

        /* =============================================
           THEME SELECTOR GRID
           ============================================= */
        .theme-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-bottom: 16px; }
        .theme-card {
            background: var(--bg-base); border: 1px solid var(--border);
            border-radius: var(--radius-md); padding: 14px; cursor: pointer;
            transition: all var(--transition); text-align: left; position: relative;
        }
        .theme-card:hover { border-color: var(--accent); }
        .theme-card.active { border-color: var(--accent); background: var(--accent-subtle); }
        .theme-card h4 { font-size: 0.88em; font-weight: 700; color: var(--text-primary); margin-bottom: 3px; }
        .theme-card p { font-size: 0.76em; color: var(--text-muted); margin-bottom: 8px; }
        .theme-dots { display: flex; gap: 5px; }
        .dot { width: 9px; height: 9px; border-radius: 50%; display: inline-block; }

        /* =============================================
           STATUS BAR
           ============================================= */
        .status-bar {
            margin-top: 14px; padding: 10px 16px; border-radius: var(--radius-md);
            font-size: 0.85em; font-weight: 600; display: none; align-items: center; gap: 8px;
        }
        .status-bar.show { display: flex; }
        .status-bar.success { background: rgba(5,150,105,0.1); border: 1px solid rgba(5,150,105,0.25); color: var(--accent-green); }
        .status-bar.error   { background: rgba(220,38,38,0.1); border: 1px solid rgba(220,38,38,0.25); color: #dc2626; }
        .status-bar.info    { background: var(--accent-subtle); border: 1px solid rgba(79,70,229,0.25); color: var(--accent); }

        /* =============================================
           PREVIEW PANEL (RIGHT SIDE)
           ============================================= */
        .preview-panel { position: sticky; top: 80px; }
        .preview-toolbar {
            display: flex; align-items: center; justify-content: space-between;
            padding: 14px 20px; border-bottom: 1px solid var(--border);
            background: var(--bg-raised);
        }
        .preview-title { font-size: 0.85em; font-weight: 700; color: var(--text-secondary); display: flex; align-items: center; gap: 8px; }
        .status-dot { width: 8px; height: 8px; border-radius: 50%; background: var(--accent-green); }
        .preview-actions { display: flex; gap: 8px; align-items: center; }

        .preview-frame {
            height: calc(100vh - 210px); min-height: 540px;
            width: 100%; border: none; background: #fff; display: block;
        }
        .preview-placeholder {
            height: calc(100vh - 210px); min-height: 540px;
            display: flex; flex-direction: column; align-items: center; justify-content: center;
            gap: 14px; color: var(--text-muted); padding: 40px; text-align: center;
        }
        .preview-placeholder i { font-size: 3em; opacity: 0.35; color: var(--accent); }

        /* =============================================
           DOWNLOAD DROPDOWN MENU
           ============================================= */
        .dropdown-wrapper { position: relative; }
        .dropdown-menu {
            display: none; position: absolute; right: 0; top: calc(100% + 6px);
            background: var(--bg-surface); border: 1px solid var(--border);
            border-radius: var(--radius-md); min-width: 200px;
            box-shadow: var(--shadow-lg); z-index: 200; overflow: hidden;
        }
        .dropdown-menu.open { display: block; animation: dropDown 0.15s ease; }
        @keyframes dropDown { from { opacity: 0; transform: translateY(-4px); } to { opacity: 1; transform: translateY(0); } }
        .dropdown-item {
            display: flex; align-items: center; gap: 10px; padding: 11px 16px;
            font-size: 0.85em; font-weight: 600; color: var(--text-primary); cursor: pointer;
            border-bottom: 1px solid var(--border); transition: background var(--transition);
        }
        .dropdown-item:last-child { border-bottom: none; }
        .dropdown-item:hover { background: var(--accent-subtle); }
        .dropdown-item .fmt { font-size: 0.75em; color: var(--text-muted); margin-left: auto; }

        /* =============================================
           FOOTER
           ============================================= */
        .footer-strip {
            border-top: 1px solid var(--border); padding: 20px 24px;
            display: flex; align-items: center; justify-content: center; gap: 32px; flex-wrap: wrap;
            background: var(--bg-surface); margin-top: auto;
        }
        .feature-chip { display: flex; align-items: center; gap: 8px; font-size: 0.8em; color: var(--text-secondary); }
        .feature-chip i { color: var(--accent); }

        .spin { animation: spin 0.8s linear infinite; }
        @keyframes spin { to { transform: rotate(360deg); } }
    </style>
</head>
<body>
    <div class="bg-grid"></div>

    <!-- WELCOME MODAL -->
    <div class="modal-overlay" id="welcomeModal">
        <div class="modal-card">
            <div class="modal-icon"><i class="fas fa-file-invoice"></i></div>
            <h2>Welcome to Portfolio Studio</h2>
            <p>Convert your resume text into a clean, editable single-page web portfolio. Have you built a portfolio with us before?</p>

            <div class="modal-options">
                <div class="modal-option-btn" onclick="handleWelcomeChoice('new')">
                    <i class="fas fa-user-plus"></i>
                    <div class="opt-title">First Time User</div>
                    <div class="opt-desc">Create a new portfolio from scratch or guided sample</div>
                </div>
                <div class="modal-option-btn" onclick="handleWelcomeChoice('returning')">
                    <i class="fas fa-rotate-left"></i>
                    <div class="opt-title">Returning User</div>
                    <div class="opt-desc">Load & edit your saved resume data instantly</div>
                </div>
            </div>

            <div class="modal-close-link" onclick="closeModal()">Skip to Editor →</div>
        </div>
    </div>

    <div class="app-root">
        <!-- NAVBAR -->
        <nav class="topbar">
            <a href="/" class="topbar-brand">
                <div class="brand-icon">P</div>
                <span class="brand-name">Portfolio<span>Studio</span></span>
            </a>
            <div class="topbar-right">
                <div class="save-indicator" id="saveIndicator"><i class="fas fa-check"></i> Saved</div>
                <button class="btn btn-secondary btn-sm" onclick="toggleDashboardTheme()" id="themeToggleBtn" title="Toggle light / dark mode"><i class="fas fa-moon" id="themeIcon"></i> <span id="themeLabel">Dark</span></button>
                <button class="btn btn-secondary btn-sm" onclick="openModal()"><i class="fas fa-user-circle"></i> Welcome</button>
                <button class="btn btn-primary btn-sm" onclick="document.getElementById('resumeFileInput').click()"><i class="fas fa-file-upload"></i> Upload Resume</button>
                <input type="file" id="resumeFileInput" style="display:none" accept=".txt,.pdf,.docx,.doc" onchange="if(this.files[0]) uploadResumeFile(this.files[0])">
                <a href="/portfolio.html" target="_blank" class="btn btn-secondary btn-sm"><i class="fas fa-external-link-alt"></i> Full View</a>
            </div>
        </nav>

        <!-- HERO HEADER -->
        <header class="hero">
            <h1>Transform Your Resume into a <span class="accent-text">Web Portfolio</span></h1>
            <p class="hero-sub">Fill out your resume details below or load the sample text to generate your customizable portfolio site.</p>
        </header>

        <!-- WORKSPACE (INPUT & PREVIEW) -->
        <main class="workspace">

            <!-- LEFT SIDE: STEPPED INPUT FORM PANEL -->
            <div class="panel">
                <div class="panel-header">
                    <div class="panel-title"><i class="fas fa-edit"></i> Resume Inputs</div>
                    <div style="display:flex; gap:8px;">
                        <button class="btn btn-amber btn-sm" onclick="loadSampleResume()"><i class="fas fa-download"></i> Load Sample</button>
                        <button class="btn btn-secondary btn-sm" onclick="resetForm()"><i class="fas fa-trash"></i> Reset</button>
                    </div>
                </div>

                <div class="panel-body">
                    <!-- SCORE BAR -->
                    <div class="score-card">
                        <div class="score-info">
                            <div class="score-title">
                                <span>Portfolio Completeness</span>
                                <span id="scoreText">0%</span>
                            </div>
                            <div class="score-bar-bg">
                                <div class="score-bar-fill" id="scoreFill"></div>
                            </div>
                        </div>
                    </div>

                    <!-- TAB NAVIGATION -->
                    <div class="form-tabs">
                        <button class="tab-btn active" onclick="switchTab('tab1', this)"><i class="fas fa-user"></i> 1. Profile</button>
                        <button class="tab-btn" onclick="switchTab('tab2', this)"><i class="fas fa-code"></i> 2. Skills & Exp</button>
                        <button class="tab-btn" onclick="switchTab('tab3', this)"><i class="fas fa-folder-open"></i> 3. Projects</button>
                        <button class="tab-btn" onclick="switchTab('tab4', this)"><i class="fas fa-palette"></i> 4. Layout Style</button>
                    </div>

                    <form id="portfolioForm" onsubmit="event.preventDefault(); generatePortfolio();">
                        <textarea id="resumeInput" style="display:none;"></textarea>

                        <!-- TAB 1: PROFILE & SUMMARY -->
                        <div class="tab-content active" id="tab1">
                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-id-card"></i> Full Name</label>
                                </div>
                                <input id="f_name" class="field-input" type="text" placeholder="e.g. Nandini Saraswat" oninput="onFieldChange()">
                            </div>

                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-briefcase"></i> Professional Title</label>
                                </div>
                                <input id="f_title" class="field-input" type="text" placeholder="e.g. Computer Science Student | Web Development Intern" oninput="onFieldChange()">
                            </div>

                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-align-left"></i> Summary</label>
                                    <button type="button" class="btn-ai-enhance" onclick="aiEnhance('summary')"><i class="fas fa-magic"></i> Auto-Format</button>
                                </div>
                                <textarea id="f_summary" class="field-input" rows="4" placeholder="Brief summary of your technical background and career goals..." oninput="onFieldChange()"></textarea>
                            </div>

                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-envelope"></i> Email Address</label>
                                </div>
                                <input id="f_email" class="field-input" type="email" placeholder="contact@example.com" oninput="onFieldChange()">
                            </div>

                            <div style="display:grid; grid-template-columns: 1fr 1fr; gap:14px;">
                                <div class="field-group">
                                    <label class="field-label"><i class="fab fa-linkedin"></i> LinkedIn</label>
                                    <input id="f_linkedin" class="field-input" type="text" placeholder="linkedin.com/in/example" oninput="onFieldChange()">
                                </div>
                                <div class="field-group">
                                    <label class="field-label"><i class="fab fa-github"></i> GitHub</label>
                                    <input id="f_github" class="field-input" type="text" placeholder="github.com/example" oninput="onFieldChange()">
                                </div>
                            </div>
                        </div>

                        <!-- TAB 2: SKILLS & EXPERIENCE -->
                        <div class="tab-content" id="tab2">
                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-tools"></i> Skills (Comma Separated)</label>
                                </div>
                                <input id="f_skills" class="field-input" type="text" placeholder="Python, Java, SQL, HTML/CSS, JavaScript, React, Flask, Git" oninput="onFieldChange()">
                                <div class="field-hint-text">Separate skills with commas</div>
                            </div>

                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-building"></i> Work Experience</label>
                                    <button type="button" class="btn-ai-enhance" onclick="aiEnhance('experience')"><i class="fas fa-magic"></i> Auto-Format</button>
                                </div>
                                <textarea id="f_experience" class="field-input" rows="5" placeholder="Format: Role | Company | Year | Responsibilities
e.g. Web Development Intern | Infosys | 2025 | Developed responsive user interfaces" oninput="onFieldChange()"></textarea>
                                <div class="field-hint-text">Format per line: Role | Company | Year | Responsibilities</div>
                            </div>
                        </div>

                        <!-- TAB 3: PROJECTS & EDUCATION -->
                        <div class="tab-content" id="tab3">
                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-layer-group"></i> Projects</label>
                                    <button type="button" class="btn-ai-enhance" onclick="aiEnhance('projects')"><i class="fas fa-magic"></i> Auto-Format</button>
                                </div>
                                <textarea id="f_projects" class="field-input" rows="5" placeholder="Format: Title | Description | Technologies
e.g. Zenith 2 | Project management platform | Python, Flask, SQL" oninput="onFieldChange()"></textarea>
                            </div>

                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-graduation-cap"></i> Education</label>
                                </div>
                                <textarea id="f_education" class="field-input" rows="3" placeholder="Format: Degree | Institution | Year
e.g. B.Tech Computer Science | GLA University | 2025 - Present" oninput="onFieldChange()"></textarea>
                            </div>
                        </div>

                        <!-- TAB 4: ACHIEVEMENTS & THEME -->
                        <div class="tab-content" id="tab4">
                            <div class="field-group">
                                <div class="field-label-row">
                                    <label class="field-label"><i class="fas fa-trophy"></i> Achievements & Leadership</label>
                                </div>
                                <textarea id="f_achievements" class="field-input" rows="3" placeholder="One entry per line...
e.g. Overall Coordinator, College Hackathon (2025)" oninput="onFieldChange()"></textarea>
                            </div>

                            <div class="field-group">
                                <label class="field-label" style="margin-bottom:12px;"><i class="fas fa-palette"></i> Portfolio Design Theme</label>
                                <div class="theme-grid">
                                    <div class="theme-card active" onclick="selectTheme('4', this)">
                                        <h4>💼 Minimalist</h4>
                                        <p>Clean & Corporate Light (Default)</p>
                                        <div class="theme-dots"><span class="dot" style="background:#f8fafc"></span><span class="dot" style="background:#3b82f6"></span><span class="dot" style="background:#0f172a"></span></div>
                                    </div>
                                    <div class="theme-card" onclick="selectTheme('2', this)">
                                        <h4>📰 Editorial</h4>
                                        <p>Clean Magazine Serif</p>
                                        <div class="theme-dots"><span class="dot" style="background:#ffffff"></span><span class="dot" style="background:#dc2626"></span><span class="dot" style="background:#1c1917"></span></div>
                                    </div>
                                    <div class="theme-card" onclick="selectTheme('1', this)">
                                        <h4>✨ Glassmorphism</h4>
                                        <p>Vibrant Gradient</p>
                                        <div class="theme-dots"><span class="dot" style="background:#a855f7"></span><span class="dot" style="background:#38bdf8"></span><span class="dot" style="background:#0f0c29"></span></div>
                                    </div>
                                    <div class="theme-card" onclick="selectTheme('3', this)">
                                        <h4>💻 Terminal</h4>
                                        <p>Developer Monospace</p>
                                        <div class="theme-dots"><span class="dot" style="background:#0d1117"></span><span class="dot" style="background:#58a6ff"></span><span class="dot" style="background:#7ee787"></span></div>
                                    </div>
                                </div>
                            </div>
                        </div>

                        <!-- SUBMIT BUTTON -->
                        <div style="margin-top: 20px;">
                            <button type="submit" id="generateBtn" class="btn btn-primary btn-full"><i class="fas fa-globe"></i> Generate & View Portfolio</button>
                        </div>
                    </form>

                    <div id="statusBar" class="status-bar"></div>
                </div>
            </div>

            <!-- RIGHT SIDE: LIVE PREVIEW PANEL -->
            <div class="panel preview-panel">
                <div class="preview-toolbar">
                    <div class="preview-title">
                        <div class="status-dot"></div> Live Webpage Preview
                    </div>
                    <div class="preview-actions">
                        <div class="dropdown-wrapper">
                            <button class="btn btn-green btn-sm" onclick="toggleDownloadDropdown()"><i class="fas fa-download"></i> Download <i class="fas fa-chevron-down" style="font-size:0.7em"></i></button>
                            <div class="dropdown-menu" id="downloadDropdown">
                                <div class="dropdown-item" onclick="downloadAs('html')"><i class="fas fa-code" style="color:#3b82f6"></i> Download HTML <span class="fmt">.html</span></div>
                                <div class="dropdown-item" onclick="downloadAs('pdf')"><i class="fas fa-file-pdf" style="color:#ef4444"></i> Download PDF <span class="fmt">.pdf</span></div>
                                <div class="dropdown-item" onclick="downloadAs('docx')"><i class="fas fa-file-word" style="color:#2563eb"></i> Download Word <span class="fmt">.doc</span></div>
                            </div>
                        </div>
                        <a href="/portfolio.html" target="_blank" class="btn btn-secondary btn-sm"><i class="fas fa-expand"></i> Full Screen</a>
                    </div>
                </div>

                <div id="previewPlaceholder" class="preview-placeholder">
                    <i class="fas fa-desktop"></i>
                    <h3>Your Web Portfolio Preview Will Appear Here</h3>
                    <p>Fill out your details on the left or click "Load Sample", then click "Generate & View Portfolio".</p>
                </div>

                <iframe id="portfolioPreview" class="preview-frame" style="display:none;" src="about:blank"></iframe>
            </div>
        </main>

        <!-- FOOTER -->
        <footer class="footer-strip">
            <div class="feature-chip"><i class="fas fa-save"></i> Local Storage Saved</div>
            <div class="feature-chip"><i class="fas fa-code"></i> Python Backend Server</div>
            <div class="feature-chip"><i class="fas fa-file-download"></i> HTML / PDF / Word Export</div>
        </footer>
    </div>

    <script>
        let selectedThemeChoice = '4';

        // --- DASHBOARD THEME TOGGLE (LIGHT / DARK) ---
        function initDashboardTheme() {
            const savedTheme = localStorage.getItem('portfolio_ai_dashboard_theme') || 'light';
            setDashboardTheme(savedTheme);
        }

        function setDashboardTheme(theme) {
            document.documentElement.setAttribute('data-theme', theme);
            localStorage.setItem('portfolio_ai_dashboard_theme', theme);
            const icon = document.getElementById('themeIcon');
            const label = document.getElementById('themeLabel');
            if (icon) icon.className = theme === 'dark' ? 'fas fa-sun' : 'fas fa-moon';
            if (label) label.textContent = theme === 'dark' ? 'Light' : 'Dark';
        }

        function toggleDashboardTheme() {
            const current = document.documentElement.getAttribute('data-theme') || 'light';
            const next = current === 'dark' ? 'light' : 'dark';
            setDashboardTheme(next);
        }

        // --- WELCOME MODAL LOGIC ---
        window.addEventListener('DOMContentLoaded', () => {
            initDashboardTheme();
            const hasVisited = localStorage.getItem('portfolio_ai_visited');
            const savedData = localStorage.getItem('portfolio_ai_resume_data');
            
            if (!hasVisited) {
                openModal();
            } else if (savedData) {
                loadSavedData(JSON.parse(savedData));
            }
        });

        function openModal() {
            document.getElementById('welcomeModal').classList.add('active');
        }

        function closeModal() {
            document.getElementById('welcomeModal').classList.remove('active');
            localStorage.setItem('portfolio_ai_visited', 'true');
        }

        function handleWelcomeChoice(choice) {
            closeModal();
            if (choice === 'returning') {
                const savedData = localStorage.getItem('portfolio_ai_resume_data');
                if (savedData) {
                    loadSavedData(JSON.parse(savedData));
                    showStatus('success', 'Loaded your saved resume data!');
                    generatePortfolio();
                } else {
                    loadSampleResume();
                }
            } else {
                loadSampleResume();
            }
        }

        // --- TAB NAVIGATION LOGIC ---
        function switchTab(tabId, el) {
            document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
            document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
            el.classList.add('active');
            document.getElementById(tabId).classList.add('active');
        }

        // --- FIELD PERSISTENCE & SCORE METER ---
        function onFieldChange() {
            updateScore();
            autoSaveData();
        }

        function updateScore() {
            const fields = ['f_name', 'f_title', 'f_summary', 'f_skills', 'f_experience', 'f_education', 'f_projects', 'f_email'];
            let filledCount = 0;
            fields.forEach(id => {
                const el = document.getElementById(id);
                if (el && el.value.trim().length > 0) filledCount++;
            });
            const percent = Math.round((filledCount / fields.length) * 100);
            document.getElementById('scoreText').innerText = percent + '%';
            document.getElementById('scoreFill').style.width = percent + '%';
        }

        function autoSaveData() {
            const data = getFormDataObj();
            localStorage.setItem('portfolio_ai_resume_data', JSON.stringify(data));
            const indicator = document.getElementById('saveIndicator');
            if (indicator) {
                indicator.style.opacity = '1';
                setTimeout(() => { indicator.style.opacity = '0.7'; }, 1000);
            }
        }

        function getFormDataObj() {
            return {
                name: document.getElementById('f_name').value,
                title: document.getElementById('f_title').value,
                summary: document.getElementById('f_summary').value,
                skills: document.getElementById('f_skills').value,
                experience: document.getElementById('f_experience').value,
                education: document.getElementById('f_education').value,
                projects: document.getElementById('f_projects').value,
                achievements: document.getElementById('f_achievements').value,
                email: document.getElementById('f_email').value,
                linkedin: document.getElementById('f_linkedin').value,
                github: document.getElementById('f_github').value,
                theme: selectedThemeChoice
            };
        }

        function loadSavedData(data) {
            if (!data) return;
            document.getElementById('f_name').value = data.name || '';
            document.getElementById('f_title').value = data.title || '';
            document.getElementById('f_summary').value = data.summary || '';
            document.getElementById('f_skills').value = data.skills || '';
            document.getElementById('f_experience').value = data.experience || '';
            document.getElementById('f_education').value = data.education || '';
            document.getElementById('f_projects').value = data.projects || '';
            document.getElementById('f_achievements').value = data.achievements || '';
            document.getElementById('f_email').value = data.email || '';
            document.getElementById('f_linkedin').value = data.linkedin || '';
            document.getElementById('f_github').value = data.github || '';
            if (data.theme) {
                selectedThemeChoice = data.theme;
            }
            updateScore();
        }

        function resetForm() {
            if (confirm("Are you sure you want to reset all input fields?")) {
                document.getElementById('portfolioForm').reset();
                localStorage.removeItem('portfolio_ai_resume_data');
                updateScore();
                showStatus('info', 'Form cleared.');
            }
        }

        // --- AUTO-FORMAT HELPER ---
        function aiEnhance(field) {
            if (field === 'summary') {
                const cur = document.getElementById('f_summary').value;
                if (!cur.trim()) {
                    document.getElementById('f_summary').value = "Motivated and detail-oriented professional with strong technical foundations in computer science, software design, and problem solving. Proven track record of delivering web applications and collaborating on team projects.";
                } else {
                    document.getElementById('f_summary').value = "Driven software specialist specializing in full-stack web solutions and data-driven systems. " + cur.trim();
                }
            } else if (field === 'experience') {
                const cur = document.getElementById('f_experience').value;
                if (!cur.trim()) {
                    document.getElementById('f_experience').value = "Software Engineering Intern | Tech Solutions | 2025 | Developed responsive user interfaces using HTML, CSS, JavaScript and optimized API endpoints.";
                }
            } else if (field === 'projects') {
                const cur = document.getElementById('f_projects').value;
                if (!cur.trim()) {
                    document.getElementById('f_projects').value = "Smart Portfolio Generator | Web app that converts raw text resumes into structured HTML portfolios using Python | Python, HTML, CSS, JavaScript";
                }
            }
            onFieldChange();
            showStatus('success', 'Form text updated!');
        }

        // --- SERIALIZE & GENERATE ---
        function serializeFields() {
            const v = id => document.getElementById(id).value.trim();
            const parts = [
                '__STRUCTURED__',
                'Name: ' + v('f_name'),
                'Title: ' + v('f_title'),
                'Summary: ' + v('f_summary').replace(/\\n/g, ' '),
                'Skills: ' + v('f_skills'),
                'Experience: ' + v('f_experience'),
                'Education: ' + v('f_education'),
                'Projects: ' + v('f_projects'),
                'Achievements: ' + v('f_achievements'),
                'Email: ' + v('f_email'),
                'LinkedIn: ' + v('f_linkedin'),
                'GitHub: ' + v('f_github')
            ];
            const result = parts.join('\\n');
            document.getElementById('resumeInput').value = result;
            return result;
        }

        async function generatePortfolio() {
            const btn = document.getElementById('generateBtn');
            const nameVal = document.getElementById('f_name').value.trim();
            if (!nameVal) {
                showStatus('error', 'Please fill in at least your Full Name before generating!');
                return;
            }

            const payloadText = serializeFields();
            btn.disabled = true;
            btn.innerHTML = '<i class="fas fa-spinner fa-spin"></i> Generating Portfolio...';
            showStatus('info', 'Building portfolio webpage...');

            try {
                const res = await fetch('/api/generate', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ resume_text: payloadText, theme_choice: selectedThemeChoice })
                });
                const result = await res.json();
                if (result.success) {
                    const iframe = document.getElementById('portfolioPreview');
                    document.getElementById('previewPlaceholder').style.display = 'none';
                    iframe.style.display = 'block';
                    iframe.src = '/portfolio.html?t=' + new Date().getTime();
                    showStatus('success', '✅ Portfolio generated! Click any text in the preview to edit directly.');
                } else {
                    showStatus('error', 'Failed to generate portfolio.');
                }
            } catch (err) {
                showStatus('error', 'Could not reach server.');
            } finally {
                btn.disabled = false;
                btn.innerHTML = '<i class="fas fa-globe"></i> Generate & View Portfolio';
            }
        }

        function selectTheme(choice, el) {
            selectedThemeChoice = choice;
            document.querySelectorAll('.theme-card').forEach(c => c.classList.remove('active'));
            el.classList.add('active');
            onFieldChange();
            generatePortfolio();
        }

        // --- SAMPLE LOAD ---
        async function loadSampleResume() {
            try {
                const res = await fetch('/api/get-resume');
                const text = await res.text();
                if (text) {
                    const lines = text.split('\\n').map(l => l.trim()).filter(Boolean);
                    let name='', title='', summary='', skills='', experience=[], education=[], projects=[], achievements=[], email='', linkedin='', github='';
                    let mode = 'header', hIdx = 0;

                    for (const line of lines) {
                        const low = line.toLowerCase();
                        if (low === 'summary' || low === 'professional summary') { mode = 'summary'; continue; }
                        if (low === 'skills') { mode = 'skills'; continue; }
                        if (low === 'experience' || low === 'work experience') { mode = 'experience'; continue; }
                        if (low === 'education') { mode = 'education'; continue; }
                        if (low === 'projects') { mode = 'projects'; continue; }
                        if (low === 'achievements' || low === 'awards' || low === 'awards & leadership') { mode = 'achievements'; continue; }

                        if (mode === 'header') {
                            if (hIdx === 0) name = line;
                            else if (/@|linkedin|github|phone|\\+/i.test(line)) {
                                if (!email) email = (line.match(/[\\w.+-]+@[\\w-]+\\.[\\w.]+/)||[name])[0];
                                if (!linkedin) linkedin = (line.match(/linkedin\\.com\\/\\S+/i)||[''])[0];
                                if (!github) github = (line.match(/github\\.com\\/\\S+/i)||[''])[0];
                            } else if (!title) title = line;
                            hIdx++;
                        }
                        else if (mode === 'summary') summary += (summary ? ' ' : '') + line;
                        else if (mode === 'skills') skills += (skills ? ', ' : '') + line.replace(/^[-•*]\\s*/, '');
                        else if (mode === 'experience') experience.push(line.replace(/^[-•*]\\s*/, ''));
                        else if (mode === 'education') education.push(line.replace(/^[-•*]\\s*/, ''));
                        else if (mode === 'projects') projects.push(line.replace(/^[-•*]\\s*/, ''));
                        else if (mode === 'achievements') achievements.push(line.replace(/^[-•*]\\s*/, ''));
                    }

                    document.getElementById('f_name').value = name;
                    document.getElementById('f_title').value = title;
                    document.getElementById('f_summary').value = summary;
                    document.getElementById('f_skills').value = skills;
                    document.getElementById('f_experience').value = experience.join('\\n');
                    document.getElementById('f_education').value = education.join('\\n');
                    document.getElementById('f_projects').value = projects.join('\\n');
                    document.getElementById('f_achievements').value = achievements.join('\\n');
                    document.getElementById('f_email').value = email || 'contact@example.com';
                    document.getElementById('f_linkedin').value = linkedin || 'linkedin.com/in/example';
                    document.getElementById('f_github').value = github || 'github.com/example';

                    onFieldChange();
                    showStatus('success', 'Sample data loaded!');
                    generatePortfolio();
                }
            } catch (err) {
                showStatus('error', 'Could not load sample.');
            }
        }

        // --- RESUME FILE UPLOAD (PDF / DOCX / TXT) ---
        async function uploadResumeFile(file) {
            if (!file) return;
            showStatus('info', `Uploading & converting "${file.name}"...`);
            const reader = new FileReader();
            reader.onload = async (event) => {
                const base64Data = event.target.result;
                try {
                    const res = await fetch('/api/upload-resume', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ filename: file.name, file_data: base64Data })
                    });
                    const result = await res.json();
                    if (result.success && result.data) {
                        populateFormWithData(result.data);
                        autoSaveData();
                        showStatus('success', `✅ "${file.name}" parsed & loaded!`);
                        generatePortfolio();
                    } else {
                        showStatus('error', result.error || 'Failed to extract text from file.');
                    }
                } catch (err) {
                    showStatus('error', 'Error uploading: ' + err.message);
                }
            };
            reader.readAsDataURL(file);
        }

        function populateFormWithData(data) {
            if (!data) return;
            const NL = String.fromCharCode(10);
            document.getElementById('f_name').value = data.name || '';
            document.getElementById('f_title').value = data.headline || '';
            document.getElementById('f_summary').value = data.professional_summary || data.summary || '';
            if (Array.isArray(data.skills)) {
                document.getElementById('f_skills').value = data.skills.join(', ');
            } else { document.getElementById('f_skills').value = data.skills || ''; }
            if (Array.isArray(data.experience)) {
                document.getElementById('f_experience').value = data.experience.map(e => {
                    if (typeof e === 'string') return e;
                    const resp = Array.isArray(e.responsibilities) ? e.responsibilities.join(' | ') : (e.description || '');
                    return [e.role||e.title||'', e.company||'', e.duration||'', resp].filter(Boolean).join(' | ');
                }).join(NL);
            } else { document.getElementById('f_experience').value = data.experience || ''; }
            if (Array.isArray(data.education)) {
                document.getElementById('f_education').value = data.education.map(e => {
                    if (typeof e === 'string') return e;
                    return [e.degree||'', e.institution||'', e.duration||''].filter(Boolean).join(' | ');
                }).join(NL);
            } else { document.getElementById('f_education').value = data.education || ''; }
            if (Array.isArray(data.projects)) {
                document.getElementById('f_projects').value = data.projects.map(p => {
                    if (typeof p === 'string') return p;
                    const tech = Array.isArray(p.technologies) ? p.technologies.join(', ') : (p.technologies || '');
                    return [p.title||p.name||'', p.description||'', tech].filter(Boolean).join(' | ');
                }).join(NL);
            } else { document.getElementById('f_projects').value = data.projects || ''; }
            if (Array.isArray(data.achievements)) {
                document.getElementById('f_achievements').value = data.achievements.map(a => typeof a === 'string' ? a : (a.title || a.description || '')).join(NL);
            } else { document.getElementById('f_achievements').value = data.achievements || ''; }
            const contact = data.contact || {};
            document.getElementById('f_email').value = contact.email || '';
            document.getElementById('f_linkedin').value = contact.linkedin || '';
            document.getElementById('f_github').value = contact.github || '';
            updateScore();
        }

        function importDataJSON(e) {
            const file = e.target.files[0];
            if (!file) return;
            const reader = new FileReader();
            reader.onload = (event) => {
                try {
                    const data = JSON.parse(event.target.result);
                    loadSavedData(data);
                    autoSaveData();
                    showStatus('success', 'JSON Resume imported successfully!');
                    generatePortfolio();
                } catch (err) {
                    showStatus('error', 'Invalid JSON file.');
                }
            };
            reader.readAsText(file);
        }

        // --- DOWNLOAD DROPDOWN LOGIC ---
        function toggleDownloadDropdown() {
            document.getElementById('downloadDropdown').classList.toggle('open');
        }
        document.addEventListener('click', (e) => {
            const wrapper = document.querySelector('.dropdown-wrapper');
            if (wrapper && !wrapper.contains(e.target)) {
                document.getElementById('downloadDropdown').classList.remove('open');
            }
        });

        async function downloadAs(format) {
            document.getElementById('downloadDropdown').classList.remove('open');
            showStatus('info', 'Preparing download...');
            if (format === 'html') {
                const a = document.createElement('a');
                a.href = '/portfolio.html?t=' + new Date().getTime();
                a.download = 'portfolio.html';
                a.click();
                showStatus('success', 'HTML file downloaded!');
            } else if (format === 'pdf') {
                const printWin = window.open('/portfolio.html?t=' + new Date().getTime(), '_blank');
                printWin.addEventListener('load', () => setTimeout(() => printWin.print(), 600));
                showStatus('success', 'Print window opened. Select "Save as PDF".');
            } else if (format === 'docx') {
                const a = document.createElement('a');
                a.href = '/api/download-docx';
                a.download = 'portfolio.doc';
                a.click();
                showStatus('success', 'Word document downloaded!');
            }
        }

        function showStatus(type, msg) {
            const bar = document.getElementById('statusBar');
            bar.className = 'status-bar show ' + type;
            bar.innerText = msg;
            setTimeout(() => { bar.classList.remove('show'); }, 6000);
        }
    </script>
</body>
</html>"""

# ==========================================================
# HTTP SERVER REQUEST HANDLER
# ==========================================================
class WebAppHandler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass

    def do_GET(self):
        url_path = urllib.parse.urlparse(self.path).path

        if url_path == '/' or url_path == '/index.html':
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.end_headers()
            self.wfile.write(WEB_APP_HTML.encode('utf-8'))

        elif url_path == '/api/get-resume':
            resume_file = os.path.join(BASE_DIR, 'resume.txt')
            content = ""
            if os.path.exists(resume_file):
                with open(resume_file, 'r', encoding='utf-8') as f:
                    content = f.read()
            self.send_response(200)
            self.send_header('Content-Type', 'text/plain; charset=utf-8')
            self.end_headers()
            self.wfile.write(content.encode('utf-8'))

        elif url_path == '/api/download-docx':
            # Generate a Word-compatible .doc file from portfolio.html
            portfolio_path = os.path.join(BASE_DIR, 'portfolio.html')
            style_path = os.path.join(BASE_DIR, 'style.css')
            if os.path.exists(portfolio_path):
                with open(portfolio_path, 'r', encoding='utf-8') as f:
                    portfolio_html = f.read()
                # Read CSS to inline it
                css_content = ""
                if os.path.exists(style_path):
                    with open(style_path, 'r', encoding='utf-8') as f:
                        css_content = f.read()
                # Wrap in Word-compatible MHTML format
                word_html = f"""<html xmlns:o="urn:schemas-microsoft-com:office:office"
xmlns:w="urn:schemas-microsoft-com:office:word"
xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="UTF-8">
<meta http-equiv="Content-Type" content="text/html; charset=utf-8">
<!--[if gte mso 9]><xml><w:WordDocument><w:View>Print</w:View></w:WordDocument></xml><![endif]-->
<style>{css_content}</style>
</head>
{portfolio_html[portfolio_html.find('<body'):]}"""
                doc_bytes = word_html.encode('utf-8')
                self.send_response(200)
                self.send_header('Content-Type', 'application/msword')
                self.send_header('Content-Disposition', 'attachment; filename="portfolio.doc"')
                self.send_header('Content-Length', str(len(doc_bytes)))
                self.end_headers()
                self.wfile.write(doc_bytes)
            else:
                self.send_error(404, "portfolio.html not found. Generate it first.")

        elif url_path == '/portfolio.html' or url_path == '/style.css':
            file_name = url_path.lstrip('/')
            file_path = os.path.join(BASE_DIR, file_name)
            if os.path.exists(file_path):
                content_type = 'text/html; charset=utf-8' if file_name.endswith('.html') else 'text/css; charset=utf-8'
                self.send_response(200)
                self.send_header('Content-Type', content_type)
                self.end_headers()
                with open(file_path, 'rb') as f:
                    self.wfile.write(f.read())
            else:
                self.send_error(404, "File Not Found")
        else:
            self.send_error(404, "Not Found")

    def do_POST(self):
        if self.path == '/api/upload-resume':
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            try:
                content_type = self.headers.get('Content-Type', '')
                filename = 'uploaded_resume.txt'
                file_bytes = b''

                if 'application/json' in content_type:
                    payload = json.loads(body.decode('utf-8'))
                    filename = payload.get('filename', 'uploaded_resume.txt')
                    file_data_b64 = payload.get('file_data', '')
                    if file_data_b64:
                        if ',' in file_data_b64:
                            file_data_b64 = file_data_b64.split(',', 1)[1]
                        file_bytes = base64.b64decode(file_data_b64)
                    else:
                        file_bytes = payload.get('text', '').encode('utf-8')
                else:
                    filename = self.headers.get('X-Filename', 'uploaded_resume.txt')
                    file_bytes = body

                extracted_text = extract_text_from_file_bytes(file_bytes, filename)
                if not extracted_text.strip():
                    self.send_response(400)
                    self.send_header('Content-Type', 'application/json')
                    self.end_headers()
                    self.wfile.write(json.dumps({'success': False, 'error': 'Could not extract readable text from uploaded file.'}).encode('utf-8'))
                    return

                cleaned_resume = clean_resume_text(extracted_text)

                # Save extracted text to resume.txt
                resume_file_path = os.path.join(BASE_DIR, 'resume.txt')
                with open(resume_file_path, 'w', encoding='utf-8') as f:
                    f.write(cleaned_resume)

                # Parse portfolio data
                portfolio_data = None
                parse_source = "structured"

                if extracted_text.startswith('__STRUCTURED__'):
                    portfolio_data = parse_structured_input(extracted_text)

                if not portfolio_data or not _has_meaningful_data(portfolio_data):
                    model = configure_gemini()
                    if model:
                        prompt = build_prompt(cleaned_resume)
                        raw_resp = call_gemini(model, prompt)
                        if raw_resp:
                            portfolio_data = parse_and_validate_json(raw_resp)
                            if _has_meaningful_data(portfolio_data):
                                parse_source = "gemini"

                if not portfolio_data or not _has_meaningful_data(portfolio_data):
                    portfolio_data = parse_resume_locally(cleaned_resume)
                    parse_source = "local"

                theme_choice = self.headers.get('X-Theme-Choice', '4')
                theme_class, theme_name = THEME_MAP.get(theme_choice, THEME_MAP["4"])
                generate_html(portfolio_data, theme_class)

                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                res_payload = {
                    'success': True,
                    'text': cleaned_resume,
                    'data': portfolio_data,
                    'source': parse_source,
                    'filename': filename
                }
                self.wfile.write(json.dumps(res_payload).encode('utf-8'))

            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'success': False, 'error': str(e)}).encode('utf-8'))

        elif self.path == '/api/generate':
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            try:
                payload = json.loads(body.decode('utf-8'))
                resume_text = payload.get('resume_text', '')
                theme_choice = payload.get('theme_choice', '4')

                cleaned_resume = clean_resume_text(resume_text)
                theme_class, theme_name = THEME_MAP.get(theme_choice, THEME_MAP["4"])

                portfolio_data = None
                parse_source = "structured"

                # 1. Try structured input format first (from the labeled UI fields)
                if resume_text.startswith('__STRUCTURED__'):
                    portfolio_data = parse_structured_input(resume_text)

                # 2. Fall back to Gemini AI if not structured format
                if not portfolio_data or not _has_meaningful_data(portfolio_data):
                    parse_source = "local"
                    model = configure_gemini()
                    if model:
                        prompt = build_prompt(cleaned_resume)
                        raw_resp = call_gemini(model, prompt)
                        if raw_resp:
                            portfolio_data = parse_and_validate_json(raw_resp)
                            if _has_meaningful_data(portfolio_data):
                                parse_source = "gemini"

                # 3. Fall back to local parser
                if not portfolio_data or not _has_meaningful_data(portfolio_data):
                    portfolio_data = parse_resume_locally(cleaned_resume)
                    parse_source = "local"

                success = generate_html(portfolio_data, theme_class)

                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                response_json = json.dumps({'success': success, 'theme': theme_name, 'source': parse_source})
                self.wfile.write(response_json.encode('utf-8'))

            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'success': False, 'error': str(e)}).encode('utf-8'))

def run_web_app(port=5000):
    """Starts the Web Application Server on http://localhost:PORT"""
    server_address = ('', port)
    try:
        httpd = ThreadingHTTPServer(server_address, WebAppHandler)
    except OSError:
        port = 8000
        server_address = ('', port)
        httpd = ThreadingHTTPServer(server_address, WebAppHandler)

    app_url = f"http://localhost:{port}/"
    print("\n========================================================")
    print(f"🚀 Web Application is running live at: {app_url}")
    print("   Open this URL in your web browser to use the Web UI!")
    print("========================================================\n")

    try:
        webbrowser.open(app_url)
    except Exception:
        pass

    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\n[INFO] Web Application server stopped.")

if __name__ == "__main__":
    run_web_app()
